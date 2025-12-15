#!/usr/bin/env python3
"""
Script to convert LIABILITY_DISCLAIMER.md to .docx format using SymBIoTTemplate.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def parse_markdown_file(md_file_path):
    """Parse the markdown file and extract structured content"""
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    return content

def create_document_from_template(template_path, output_path, md_content):
    """Create a new document based on template with liability disclaimer content"""

    # Try to open template, if not available, create new document
    try:
        doc = Document(template_path)
        print(f"Using template: {template_path}")
    except Exception as e:
        print(f"Could not open template: {e}")
        print("Creating new document with similar styling...")
        doc = Document()

        # Set up document styling similar to typical SymBIoT template
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

    # Don't clear the body completely as it may remove sections
    # Instead, we'll add content to the existing document

    # Parse and add content line by line
    lines = md_content.split('\n')

    for line in lines:
        line = line.rstrip()

        # Skip horizontal rules
        if line.strip() == '---':
            doc.add_paragraph()
            continue

        # Main title (# )
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_heading(title, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Major headings (## )
        elif line.startswith('## '):
            heading = line[3:].strip()
            doc.add_heading(heading, level=2)

        # Sub-headings (### )
        elif line.startswith('### '):
            heading = line[4:].strip()
            doc.add_heading(heading, level=3)

        # Sub-sub-headings (#### )
        elif line.startswith('#### '):
            heading = line[5:].strip()
            doc.add_heading(heading, level=4)

        # Bullet points with - or *
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Handle bold text in markdown
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            try:
                p = doc.add_paragraph(text, style='List Bullet')
            except KeyError:
                p = doc.add_paragraph('• ' + text)

        # Numbered lists
        elif re.match(r'^\d+\.', line.strip()):
            number = re.match(r'^(\d+)\.', line.strip()).group(1)
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            try:
                p = doc.add_paragraph(text, style='List Number')
            except KeyError:
                p = doc.add_paragraph(f'{number}. {text}')

        # Empty lines
        elif line.strip() == '':
            # Don't add excessive blank paragraphs
            pass

        # Regular paragraphs
        elif line.strip():
            text = line.strip()
            # Handle bold text
            if '**' in text:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part:
                        p.add_run(part)
            else:
                doc.add_paragraph(text)

    # Add footer with branding
    try:
        if doc.sections:
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.text = "SymBIoT - Bringing families and technology together for better care."
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"Note: Could not add footer: {e}")

    # Save document
    doc.save(output_path)
    print(f"Document created successfully: {output_path}")

def main():
    # File paths
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(script_dir)
    docs_dir = os.path.join(project_root, 'docs')

    template_path = os.path.join(docs_dir, 'SymBIoTTemplate.docx')
    md_file_path = os.path.join(docs_dir, 'LIABILITY_DISCLAIMER.md')
    output_path = os.path.join(docs_dir, 'LIABILITY_DISCLAIMER.docx')

    print("=== SymBIoT Document Converter ===")
    print(f"Template: {template_path}")
    print(f"Source: {md_file_path}")
    print(f"Output: {output_path}")
    print()

    # Check if markdown file exists
    if not os.path.exists(md_file_path):
        print(f"Error: Markdown file not found: {md_file_path}")
        return

    # Parse markdown content
    print("Reading markdown content...")
    md_content = parse_markdown_file(md_file_path)

    # Create Word document
    print("Creating Word document...")
    create_document_from_template(template_path, output_path, md_content)

    print()
    print("[SUCCESS] Conversion completed successfully!")

if __name__ == '__main__':
    main()
